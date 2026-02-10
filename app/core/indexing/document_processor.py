import os
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from datetime import datetime
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    content: str
    metadata: Dict[str, Any]
    chunk_id: str
    document_id: str


@dataclass
class ProcessedDocument:
    document_id: str
    filename: str
    file_path: str
    file_type: str
    file_size: int
    chunks: List[DocumentChunk]
    total_chunks: int
    created_at: datetime
    content_hash: str


class DocumentProcessor:
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = ["\n\n", "\n", ". ", " "]
        
    def process_file(self, file_path: str, document_id: Optional[str] = None) -> ProcessedDocument:
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if document_id is None:
            document_id = self._generate_document_id(file_path)
        
        content = self._extract_text(file_path)
        content_hash = self._compute_hash(content)
        
        chunks = self._chunk_text(content, document_id, str(file_path))
        
        return ProcessedDocument(
            document_id=document_id,
            filename=file_path.name,
            file_path=str(file_path),
            file_type=file_path.suffix.lower(),
            file_size=file_path.stat().st_size,
            chunks=chunks,
            total_chunks=len(chunks),
            created_at=datetime.now(),
            content_hash=content_hash
        )
    
    def _extract_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == '.txt':
                return self._read_text_file(file_path)
            elif suffix == '.pdf':
                return self._read_pdf_file(file_path)
            elif suffix == '.docx':
                return self._read_docx_file(file_path)
            elif suffix == '.md':
                return self._read_text_file(file_path)
            elif suffix == '.html':
                return self._read_html_file(file_path)
            elif suffix == '.csv':
                return self._read_csv_file(file_path)
            else:
                logger.warning(f"Unsupported file type: {suffix}, attempting text read")
                return self._read_text_file(file_path)
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _read_text_file(self, file_path: Path) -> str:
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Could not decode file: {file_path}")
    
    def _read_pdf_file(self, file_path: Path) -> str:
        try:
            import PyPDF2
            text = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(f"\n--- Page {page_num + 1} ---\n{page_text}")
            return "\n".join(text)
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise
    
    def _read_docx_file(self, file_path: Path) -> str:
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
    
    def _read_html_file(self, file_path: Path) -> str:
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                for script in soup(["script", "style"]):
                    script.decompose()
                return soup.get_text(separator='\n', strip=True)
        except ImportError:
            logger.error("BeautifulSoup not installed. Install with: pip install beautifulsoup4")
            raise
    
    def _read_csv_file(self, file_path: Path) -> str:
        try:
            import csv
            rows = []
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(" | ".join(row))
            return "\n".join(rows)
        except Exception as e:
            logger.error(f"Error reading CSV: {str(e)}")
            raise
    
    def _chunk_text(self, text: str, document_id: str, source_path: str) -> List[DocumentChunk]:
        chunks = []
        text = self._preprocess_text(text)
        
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            
            if end < len(text):
                end = self._find_chunk_boundary(text, start, end)
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk_id = f"{document_id}_{chunk_index}"
                metadata = {
                    'source': source_path,
                    'document_id': document_id,
                    'chunk_index': chunk_index,
                    'start_char': start,
                    'end_char': end,
                    'char_count': len(chunk_text),
                    'created_at': datetime.now().isoformat()
                }
                
                chunks.append(DocumentChunk(
                    content=chunk_text,
                    metadata=metadata,
                    chunk_id=chunk_id,
                    document_id=document_id
                ))
                chunk_index += 1
            
            start = end - self.chunk_overlap if end < len(text) else end
        
        return chunks
    
    def _preprocess_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        return text.strip()
    
    def _find_chunk_boundary(self, text: str, start: int, end: int) -> int:
        for separator in self.separators:
            pos = text.rfind(separator, start, end)
            if pos > start + self.chunk_size // 2:
                return pos + len(separator)
        return end
    
    def _generate_document_id(self, file_path: Path) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        name_hash = hashlib.md5(file_path.name.encode()).hexdigest()[:8]
        return f"{timestamp}_{name_hash}"
    
    def _compute_hash(self, content: str) -> str:
        return hashlib.sha256(content.encode()).hexdigest()[:16]
